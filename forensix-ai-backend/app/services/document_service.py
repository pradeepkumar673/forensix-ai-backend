"""
app/services/document_service.py
---------------------------------
Forensic document processing service.

Responsibilities:
  • Extract raw text from PDF files (text-native and OCR fallback for scans)
  • Clean and normalise extracted text for downstream LLM processing
  • Detect and segment document sections (autopsy, toxicology, statements, etc.)
  • Parse witness / suspect statements into structured dicts
  • Chunk long documents for embedding / vector-store ingestion
  • Extract embedded images from PDFs for vision analysis
  • Compute SHA-256 checksums for evidence integrity verification

Libraries used (all free / open-source):
  • PyMuPDF  (fitz)  — primary PDF engine; fast, handles scans via integrated OCR
  • pdfplumber       — layout-aware text + table extraction
  • Pillow           — image normalisation
  • pytesseract      — OCR fallback when PyMuPDF OCR is unavailable
  • langdetect       — language detection for multilingual evidence
  • re / unicodedata — text normalisation
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Optional

import fitz  # PyMuPDF
import pdfplumber
from PIL import Image
import docx

logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------- #
# Constants                                                                     #
# --------------------------------------------------------------------------- #

# Minimum character yield per page before we fall back to OCR
_OCR_FALLBACK_THRESHOLD = 80

# Hard cap for a single LLM context window chunk (in characters)
_DEFAULT_CHUNK_SIZE   = 6_000
_DEFAULT_CHUNK_OVERLAP = 400

# Forensic section heading patterns — ordered most-specific first
_SECTION_PATTERNS: dict[str, re.Pattern[str]] = {
    "cause_of_death":       re.compile(
        r"(cause\s+of\s+death|manner\s+of\s+death|mode\s+of\s+death)", re.I),
    "external_examination": re.compile(
        r"(external\s+(examination|exam)|body\s+surface\s+examination)", re.I),
    "internal_examination": re.compile(
        r"(internal\s+(examination|exam)|organ\s+dissection|evisceration)", re.I),
    "toxicology":           re.compile(
        r"(toxicolog|drug\s+screen|substance\s+analysis|blood\s+alcohol)", re.I),
    "microscopy":           re.compile(
        r"(histolog|microscop|neuropatholog|cytolog)", re.I),
    "radiology":            re.compile(
        r"(radiolog|x[\-\s]?ray|ct\s+scan|mri|imaging\s+findings?)", re.I),
    "wounds_injuries":      re.compile(
        r"(wound|injur|lacerat|contusion|abrasion|fracture|haemorrhage|hemorrhage)", re.I),
    "identification":       re.compile(
        r"(identification|decedent\s+info|personal\s+data|identifying\s+marks?)", re.I),
    "opinion_conclusions":  re.compile(
        r"(opinion|conclusion|summary\s+of\s+findings?|pathologist.{0,20}opinion)", re.I),
    "witness_statement":    re.compile(
        r"(witness\s+statement|sworn\s+statement|declaration|affidavit|deponent)", re.I),
    "police_narrative":     re.compile(
        r"(police\s+report|incident\s+report|officer.{0,20}narrative|responding\s+officer)", re.I),
}

# Patterns for personal identifiers — used in anonymisation
_PII_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("NHS_NUMBER",    re.compile(r"\b\d{3}[\s-]\d{3}[\s-]\d{4}\b")),
    ("UK_POSTCODE",   re.compile(r"\b[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2}\b", re.I)),
    ("US_SSN",        re.compile(r"\b\d{3}-\d{2}-\d{4}\b")),
    ("PHONE",         re.compile(r"\b(\+?\d[\d\s\-().]{7,}\d)\b")),
    ("DATE_OF_BIRTH", re.compile(r"\b(D\.?O\.?B\.?|date\s+of\s+birth)\s*:?\s*[\d/\-\.]+", re.I)),
]


# --------------------------------------------------------------------------- #
# Data models                                                                   #
# --------------------------------------------------------------------------- #

class DocumentType(str, Enum):
    AUTOPSY_REPORT    = "autopsy_report"
    TOXICOLOGY_REPORT = "toxicology_report"
    WITNESS_STATEMENT = "witness_statement"
    POLICE_REPORT     = "police_report"
    MEDICAL_RECORD    = "medical_record"
    LAB_RESULT        = "lab_result"
    UNKNOWN           = "unknown"


@dataclass
class PageContent:
    """Extracted content from a single PDF page."""
    page_number:  int
    raw_text:     str
    clean_text:   str
    char_count:   int
    word_count:   int
    is_ocr:       bool          = False
    tables:       list[list]    = field(default_factory=list)
    image_count:  int           = 0
    section_tags: list[str]     = field(default_factory=list)


@dataclass
class DocumentSection:
    """A contiguous section of a document belonging to a single forensic category."""
    section_type: str
    heading:      str
    pages:        list[int]
    text:         str
    char_count:   int = 0

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class ParsedStatement:
    """
    A single processed witness / suspect statement.
    Ready for insertion into detect_contradictions().
    """
    source:         str          # "Witness A — Jane Smith, neighbour"
    role:           str          # "witness" | "suspect" | "officer" | "expert"
    text:           str          # Full cleaned statement text
    date_given:     Optional[str] = None
    location_given: Optional[str] = None
    language:       str          = "en"
    word_count:     int          = 0
    key_claims:     list[str]    = field(default_factory=list)

    def __post_init__(self):
        self.word_count = len(self.text.split())


@dataclass
class ExtractionResult:
    """
    Complete result of processing a single PDF evidence file.
    """
    file_path:      str
    file_name:      str
    checksum_sha256: str
    document_type:  DocumentType
    page_count:     int
    total_chars:    int
    total_words:    int
    full_text:      str                    # Entire document as one string
    pages:          list[PageContent]
    sections:       dict[str, DocumentSection]
    statements:     list[ParsedStatement]
    chunks:         list[str]              # LLM-ready text chunks
    embedded_images: list[bytes]           # Raw PNG bytes per embedded image
    metadata:       dict[str, Any]
    extraction_quality: str               # "good" | "partial" | "ocr_fallback" | "failed"
    warnings:       list[str]             = field(default_factory=list)
    processed_at:   str                   = field(
        default_factory=lambda: datetime.utcnow().isoformat()
    )


# --------------------------------------------------------------------------- #
# Internal utilities                                                            #
# --------------------------------------------------------------------------- #

def _sha256(path: str | Path) -> str:
    """Compute SHA-256 hex digest of a file without loading it fully into RAM."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65_536), b""):
            h.update(chunk)
    return h.hexdigest()


def _unicode_normalise(text: str) -> str:
    """Normalise unicode, collapse non-breaking spaces, fix common ligatures."""
    text = unicodedata.normalize("NFKC", text)
    # Ligatures not always resolved by NFKC
    _ligatures = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl",
                  "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
                  "\u2013": "-", "\u2014": "--"}
    for lig, rep in _ligatures.items():
        text = text.replace(lig, rep)
    return text


def _clean_text(raw: str) -> str:
    """
    Clean raw PDF-extracted text into consistent, LLM-ready prose.

    Steps:
      1. Unicode normalisation and ligature substitution
      2. Strip form-feed and null characters
      3. Collapse runs of blank lines (max 2 consecutive)
      4. Remove page headers/footers that repeat across pages (short lines
         at start/end of a block that contain only page numbers or dates)
      5. Strip trailing whitespace per line
      6. Collapse excessive internal whitespace within lines
      7. Rejoin soft-hyphenated line breaks
    """
    if not raw:
        return ""

    text = _unicode_normalise(raw)

    # Remove non-printable control characters except newline/tab
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)

    # Rejoin soft-hyphen line breaks (word- \ncontinued → wordcontinued)
    text = re.sub(r"-\s*\n\s*([a-z])", r"\1", text)

    # Strip lines that are purely page numbers or repetitive headers/footers
    lines = text.split("\n")
    cleaned_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        # Drop: pure page numbers, e.g. "- 3 -"  "Page 3 of 12"
        if re.fullmatch(r"[-–—\s]*[Pp]age\s*\d+(\s+of\s+\d+)?[-–—\s]*", stripped):
            continue
        if re.fullmatch(r"[-–—•·\s\d]+", stripped) and len(stripped) < 8:
            continue
        # Collapse excessive internal spaces
        stripped = re.sub(r"[ \t]{2,}", " ", stripped)
        cleaned_lines.append(stripped)

    text = "\n".join(cleaned_lines)

    # Collapse more than 2 consecutive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _detect_document_type(text: str, metadata: dict[str, Any]) -> DocumentType:
    """
    Heuristically determine the document type from text content and PDF metadata.
    """
    lower = text.lower()
    title = str(metadata.get("title", "")).lower()
    combined = lower[:3000] + " " + title

    scores: dict[DocumentType, int] = {t: 0 for t in DocumentType}

    # Autopsy
    if any(kw in combined for kw in
           ["autopsy", "post-mortem", "post mortem", "patholog", "cause of death",
            "manner of death", "decedent", "evisceration"]):
        scores[DocumentType.AUTOPSY_REPORT] += 3

    # Toxicology
    if any(kw in combined for kw in
           ["toxicolog", "blood alcohol", "drug screen", "urine screen",
            "vitreous", "ng/ml", "mg/l", "gas chromatography"]):
        scores[DocumentType.TOXICOLOGY_REPORT] += 3

    # Witness / statement
    if any(kw in combined for kw in
           ["witness statement", "sworn statement", "deponent", "affidavit",
            "i was", "i saw", "i heard", "i observed"]):
        scores[DocumentType.WITNESS_STATEMENT] += 3

    # Police report
    if any(kw in combined for kw in
           ["police report", "incident report", "officer", "dispatch", "responding",
            "case number", "crime scene", "arrest"]):
        scores[DocumentType.POLICE_REPORT] += 3

    # Medical record
    if any(kw in combined for kw in
           ["medical record", "patient", "diagnosis", "treatment", "prescription",
            "hospital", "clinic", "icd-"]):
        scores[DocumentType.MEDICAL_RECORD] += 3

    # Lab result
    if any(kw in combined for kw in
           ["laboratory", "lab result", "specimen", "assay", "reference range",
            "positive", "negative", "reactive"]):
        scores[DocumentType.LAB_RESULT] += 2

    best = max(scores, key=lambda k: scores[k])
    return best if scores[best] > 0 else DocumentType.UNKNOWN


def _tag_sections(pages: list[PageContent]) -> dict[str, DocumentSection]:
    """
    Walk pages in order and assign each paragraph to the most recently seen
    section heading.  Returns a dict keyed by section type.
    """
    sections: dict[str, list[tuple[int, str]]] = {k: [] for k in _SECTION_PATTERNS}
    current_section = "general"

    for page in pages:
        for paragraph in page.clean_text.split("\n\n"):
            para = paragraph.strip()
            if not para:
                continue
            # Check if this paragraph is a section heading
            matched = False
            for sec_type, pattern in _SECTION_PATTERNS.items():
                first_line = para.split("\n")[0]
                if pattern.search(first_line) and len(first_line) < 120:
                    current_section = sec_type
                    if sec_type not in sections:
                        sections[sec_type] = []
                    sections[sec_type].append((page.page_number, para))
                    page.section_tags.append(sec_type)
                    matched = True
                    break
            if not matched:
                if current_section not in sections:
                    sections[current_section] = []
                sections[current_section].append((page.page_number, para))

    result: dict[str, DocumentSection] = {}
    for sec_type, para_list in sections.items():
        if not para_list:
            continue
        section_pages = sorted(set(p for p, _ in para_list))
        section_text  = "\n\n".join(t for _, t in para_list)
        result[sec_type] = DocumentSection(
            section_type=sec_type,
            heading=sec_type.replace("_", " ").title(),
            pages=section_pages,
            text=section_text,
        )
    return result


def _chunk_text(
    text: str,
    chunk_size: int  = _DEFAULT_CHUNK_SIZE,
    overlap:    int  = _DEFAULT_CHUNK_OVERLAP,
) -> list[str]:
    """
    Split text into overlapping chunks suitable for LLM context windows or
    vector-store embedding.

    Splits prefer paragraph boundaries; falls back to sentence boundaries;
    hard-cuts at character limit as a last resort.
    """
    if len(text) <= chunk_size:
        return [text]

    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks: list[str] = []
    current_chunk: list[str] = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) + 2 > chunk_size and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            # Overlap: keep last N chars worth of paragraphs
            overlap_paras: list[str] = []
            overlap_len = 0
            for p in reversed(current_chunk):
                if overlap_len + len(p) <= overlap:
                    overlap_paras.insert(0, p)
                    overlap_len += len(p)
                else:
                    break
            current_chunk = overlap_paras
            current_len   = overlap_len

        current_chunk.append(para)
        current_len += len(para) + 2

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


# --------------------------------------------------------------------------- #
# Statement processing                                                          #
# --------------------------------------------------------------------------- #

_STATEMENT_HEADER_PATTERN = re.compile(
    r"(?:statement\s+(?:of|by|from)|witness\s+statement|suspect\s+interview|"
    r"sworn\s+statement|declaration\s+of)\s*[:\-]?\s*(?P<name>[A-Z][a-zA-Z\s\-\']{2,50})",
    re.I,
)

_DATE_PATTERN = re.compile(
    r"(?:dated?|date\s+(?:of\s+)?statement)\s*[:\-]?\s*"
    r"(?P<date>\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}|\d{1,2}\s+\w+\s+\d{4})",
    re.I,
)

_ROLE_KEYWORDS: dict[str, list[str]] = {
    "suspect":  ["suspect", "accused", "defendant", "charged", "arrested"],
    "officer":  ["officer", "detective", "sergeant", "inspector", "constable", "agent"],
    "expert":   ["expert", "pathologist", "toxicologist", "analyst", "forensic", "doctor", "dr."],
    "witness":  ["witness", "neighbour", "neighbor", "bystander", "caller", "informant"],
}


def _infer_role(header_text: str) -> str:
    lower = header_text.lower()
    for role, keywords in _ROLE_KEYWORDS.items():
        if any(kw in lower for kw in keywords):
            return role
    return "witness"


def _extract_key_claims(text: str, max_claims: int = 10) -> list[str]:
    """
    Extract atomic factual claims from a statement using sentence segmentation.
    A 'claim' is a sentence containing a first-person assertion or a
    specific time/place reference.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text)
    claim_patterns = [
        re.compile(r"\b(I\s+(was|saw|heard|noticed|went|left|arrived|found|called|told))", re.I),
        re.compile(r"\b(at\s+(?:approximately\s+)?\d{1,2}[:\.]?\d{0,2}\s*(?:am|pm)?)", re.I),
        re.compile(r"\b(on\s+(?:the\s+)?\d{1,2}(?:st|nd|rd|th)?\s+\w+)", re.I),
        re.compile(r"\b(he|she|they)\s+(was|were|had|went|came|left|said|told)", re.I),
    ]
    claims = []
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 20:
            continue
        if any(p.search(sent) for p in claim_patterns):
            claims.append(sent)
        if len(claims) >= max_claims:
            break
    return claims


# --------------------------------------------------------------------------- #
# Core public functions                                                         #
# --------------------------------------------------------------------------- #

async def extract_text_from_pdf(
    file_path: str | Path,
    use_ocr_fallback: bool = True,
    extract_tables: bool  = True,
    extract_images: bool  = True,
    password: str | None  = None,
) -> ExtractionResult:
    """
    Extract all content from a PDF evidence file.

    Strategy:
      1. Open with PyMuPDF for fast native text extraction.
      2. For any page yielding fewer than ``_OCR_FALLBACK_THRESHOLD`` characters,
         fall back to Tesseract OCR via PyMuPDF's built-in OCR bridge.
      3. Extract tables using pdfplumber (more accurate layout parsing).
      4. Extract embedded raster images as PNG bytes.
      5. Detect document type, tag sections, chunk text, parse statements.

    Parameters
    ----------
    file_path         : Absolute or relative path to the PDF.
    use_ocr_fallback  : Enable Tesseract OCR for scanned pages.
    extract_tables    : Extract tabular data via pdfplumber.
    extract_images    : Extract embedded raster images.
    password          : PDF decryption password if encrypted.

    Returns
    -------
    ExtractionResult  : Fully populated result dataclass.

    Raises
    ------
    FileNotFoundError : If the file does not exist.
    RuntimeError      : If the PDF cannot be opened (corrupt / wrong password).
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Evidence file not found: {path}")

    warnings: list[str] = []
    checksum = _sha256(path)

    # ── Open with PyMuPDF ─────────────────────────────────────────────── #
    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise RuntimeError(f"PyMuPDF could not open {path.name}: {exc}") from exc

    # pylint: disable=no-member
    if doc.is_encrypted:
        if password:
            if not doc.authenticate(password):
                raise RuntimeError("Incorrect PDF password supplied.")
        else:
            raise RuntimeError("PDF is encrypted; supply a password.")

    # pylint: disable=no-member
    pdf_metadata: dict[str, Any] = dict(doc.metadata) if doc.metadata else {}
    page_count   = doc.page_count

    # ── Page-level extraction ─────────────────────────────────────────── #
    pages:    list[PageContent] = []
    all_text: list[str]         = []
    embedded_images: list[bytes]= []
    ocr_used = False

    # pdfplumber for tables (open once, iterate pages)
    plumber_doc = pdfplumber.open(str(path)) if extract_tables else None

    for page_idx in range(page_count):
        fitz_page  = doc[page_idx]
        page_num   = page_idx + 1
        is_ocr_page = False

        # --- Native text extraction ---
        raw_text = fitz_page.get_text("text")

        # --- OCR fallback ---
        if use_ocr_fallback and len(raw_text.strip()) < _OCR_FALLBACK_THRESHOLD:
            try:
                # Render page to image, run Tesseract
                mat  = fitz.Matrix(2.0, 2.0)   # 2× zoom ≈ 144 DPI
                pix  = fitz_page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                import pytesseract
                ocr_text = pytesseract.image_to_string(img, config="--psm 6")
                if len(ocr_text.strip()) > len(raw_text.strip()):
                    raw_text    = ocr_text
                    is_ocr_page = True
                    ocr_used    = True
            except ImportError:
                warnings.append("pytesseract not installed — OCR fallback skipped.")
            except Exception as exc:
                warnings.append(f"OCR failed on page {page_num}: {exc}")

        clean = _clean_text(raw_text)

        # --- Table extraction ---
        tables: list[list] = []
        if extract_tables and plumber_doc:
            try:
                pb_page = plumber_doc.pages[page_idx]
                raw_tables = pb_page.extract_tables()
                for tbl in (raw_tables or []):
                    if tbl:
                        tables.append(tbl)
            except Exception as exc:
                warnings.append(f"Table extraction failed on page {page_num}: {exc}")

        # --- Embedded image extraction ---
        img_count = 0
        if extract_images:
            try:
                for img_info in fitz_page.get_images(full=True):
                    xref = img_info[0]
                    base_img = doc.extract_image(xref)
                    img_bytes = base_img.get("image", b"")
                    if len(img_bytes) > 2048:   # skip tiny icons
                        # Convert to PNG regardless of source format
                        pil_img  = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                        buf = io.BytesIO()
                        pil_img.save(buf, format="PNG", optimize=True)
                        embedded_images.append(buf.getvalue())
                        img_count += 1
            except Exception as exc:
                warnings.append(f"Image extraction failed on page {page_num}: {exc}")

        page_content = PageContent(
            page_number = page_num,
            raw_text    = raw_text,
            clean_text  = clean,
            char_count  = len(clean),
            word_count  = len(clean.split()),
            is_ocr      = is_ocr_page,
            tables      = tables,
            image_count = img_count,
        )
        pages.append(page_content)
        all_text.append(clean)

    if plumber_doc:
        plumber_doc.close()
    doc.close()

    full_text   = "\n\n".join(filter(None, all_text))
    total_chars = len(full_text)
    total_words = len(full_text.split())

    # ── Quality assessment ────────────────────────────────────────────── #
    ocr_pages   = sum(1 for p in pages if p.is_ocr)
    empty_pages = sum(1 for p in pages if p.char_count < 50)
    if total_chars < 100:
        quality = "failed"
    elif ocr_pages > page_count * 0.5:
        quality = "ocr_fallback"
    elif empty_pages > page_count * 0.3:
        quality = "partial"
    else:
        quality = "good"

    # ── Higher-level processing ───────────────────────────────────────── #
    doc_type  = _detect_document_type(full_text, pdf_metadata)
    sections  = _tag_sections(pages)
    chunks    = _chunk_text(full_text)

    statements = _parse_statements_from_text(
        full_text, doc_type
    )

    if ocr_used:
        warnings.insert(0, f"OCR used on {ocr_pages} of {page_count} pages.")

    return ExtractionResult(
        file_path        = str(path.resolve()),
        file_name        = path.name,
        checksum_sha256  = checksum,
        document_type    = doc_type,
        page_count       = page_count,
        total_chars      = total_chars,
        total_words      = total_words,
        full_text        = full_text,
        pages            = pages,
        sections         = sections,
        statements       = statements,
        chunks           = chunks,
        embedded_images  = embedded_images,
        metadata         = pdf_metadata,
        extraction_quality = quality,
        warnings         = warnings,
    )


# --------------------------------------------------------------------------- #

def clean_text(raw_text: str) -> str:
    """
    Public wrapper around the internal ``_clean_text`` function.

    Use this when you already have raw text (e.g. from an OCR engine or
    a plain-text file) and want to normalise it for LLM ingestion without
    running the full PDF pipeline.

    Parameters
    ----------
    raw_text : Any raw text string.

    Returns
    -------
    Cleaned, normalised text ready for LLM processing.
    """
    return _clean_text(raw_text)


# --------------------------------------------------------------------------- #

def process_statements(
    raw_statements: list[dict[str, str]],
    auto_extract_claims: bool = True,
) -> list[ParsedStatement]:
    """
    Process a list of raw statement dicts into structured ``ParsedStatement``
    objects ready for ``detect_contradictions()``.

    Parameters
    ----------
    raw_statements       : List of dicts with keys:
                             ``source``  — e.g. "Witness A — Jane Smith"
                             ``text``    — full statement text
                             ``role``    — (optional) "witness" | "suspect" | "officer" | "expert"
                             ``date``    — (optional) date given as string
                             ``location``— (optional) where statement was given
    auto_extract_claims  : If True, automatically extract atomic claims from
                           each statement for quick contradiction scanning.

    Returns
    -------
    List of ``ParsedStatement`` dataclass instances.
    """
    if not raw_statements:
        return []

    parsed: list[ParsedStatement] = []

    for item in raw_statements:
        source   = str(item.get("source", "Unknown")).strip()
        raw_text = str(item.get("text", "")).strip()
        role     = str(item.get("role", _infer_role(source))).lower()
        date     = item.get("date")
        location = item.get("location")

        if not raw_text:
            logger.warning("Empty statement body for source '%s' — skipping.", source)
            continue

        clean = _clean_text(raw_text)

        # Language detection (best-effort; requires langdetect if installed)
        language = "en"
        try:
            from langdetect import detect
            language = detect(clean[:2000]) or "en"
        except Exception:
            pass

        key_claims: list[str] = []
        if auto_extract_claims:
            key_claims = _extract_key_claims(clean)

        parsed.append(ParsedStatement(
            source         = source,
            role           = role,
            text           = clean,
            date_given     = str(date) if date else None,
            location_given = str(location) if location else None,
            language       = language,
            key_claims     = key_claims,
        ))

    logger.info("Processed %d statements (%d skipped empty).",
                len(parsed), len(raw_statements) - len(parsed))
    return parsed


# --------------------------------------------------------------------------- #

def _parse_statements_from_text(
    text: str,
    doc_type: DocumentType,
) -> list[ParsedStatement]:
    """
    Internal: attempt to auto-detect and split multiple statements embedded
    inside a single document (e.g. a combined witness interview transcript).
    """
    if doc_type not in (DocumentType.WITNESS_STATEMENT, DocumentType.POLICE_REPORT):
        return []

    # Split on patterns like "Statement of John Doe" / "WITNESS STATEMENT:"
    segments = re.split(
        r"(?=(?:statement\s+(?:of|by|from)|witness\s+statement|sworn\s+statement)"
        r"\s*[:\-]?\s*[A-Z])",
        text,
        flags=re.I,
    )

    raw_stmts: list[dict[str, str]] = []
    for seg in segments:
        seg = seg.strip()
        if len(seg) < 150:
            continue
        header_match = _STATEMENT_HEADER_PATTERN.search(seg[:300])
        source = header_match.group("name").strip() if header_match else "Unknown Witness"
        date_match = _DATE_PATTERN.search(seg[:500])
        date = date_match.group("date").strip() if date_match else None
        raw_stmts.append({"source": source, "text": seg, "date": date})

    return process_statements(raw_stmts, auto_extract_claims=True)


# --------------------------------------------------------------------------- #

def anonymise_text(text: str, replacement_map: dict[str, str] | None = None) -> str:
    """
    Redact PII from extracted text.

    Parameters
    ----------
    text             : Input text.
    replacement_map  : Optional dict mapping real names/IDs to pseudonyms,
                       e.g. {"John Smith": "SUBJECT_A"}.

    Returns
    -------
    Redacted text with PII replaced by labelled placeholders.
    """
    # Apply explicit name/ID replacements first
    if replacement_map:
        for real, pseudo in replacement_map.items():
            text = text.replace(real, pseudo)

    # Apply regex patterns
    for label, pattern in _PII_PATTERNS:
        text = pattern.sub(f"[{label}_REDACTED]", text)

    return text


# --------------------------------------------------------------------------- #

def get_section_text(result: ExtractionResult, section_type: str) -> str | None:
    """
    Convenience accessor — returns the text of a named section or None.

    Example:
        tox_text = get_section_text(result, "toxicology")
    """
    section = result.sections.get(section_type)
    return section.text if section else None

def extract_text_from_docx(file_path: str | Path) -> str:
    """Return plain text from a .docx file."""
    import docx
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")
    doc = docx.Document(str(path))
    return "\n".join(para.text for para in doc.paragraphs if para.text)


def get_tables_as_text(result: ExtractionResult) -> str:
    """
    Render all extracted tables as simple plain-text grids for LLM ingestion.
    """
    parts: list[str] = []
    for page in result.pages:
        for tbl_idx, table in enumerate(page.tables, start=1):
            header    = f"[Table {tbl_idx} — Page {page.page_number}]"
            table_str = "\n".join(
                " | ".join(str(cell or "") for cell in row)
                for row in table
            )
            parts.append(f"{header}\n{table_str}")
    return "\n\n".join(parts)


# --------------------------------------------------------------------------- #

async def extract_medical_entities(text: str) -> list[dict]:
    """
    Extract medical entities from raw text using BioClinicalBERT (if advanced vision enabled).
    """
    from app.core.config import get_settings
    if not get_settings().ENABLE_ADVANCED_VISION:
        return []
        
    try:
        from app.utils.hf_models import load_med_ner
        bundle = await load_med_ner()
        # The pipeline returns a list of dicts: [{'entity_group': 'PROBLEM', 'word': '...', 'score': 0.97}, ...]
        entities = bundle.pipeline(text)
        return entities
    except Exception as e:
        logger.warning("Medical NER extraction failed: %s", e)
        return []